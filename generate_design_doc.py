# -*- coding: utf-8 -*-
"""
BMS 博客管理系统 - 详细设计说明书生成脚本
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 颜色定义
PRIMARY_COLOR = RGBColor(0, 82, 147)
SECONDARY_COLOR = RGBColor(51, 102, 153)
ACCENT_COLOR = RGBColor(0, 112, 192)
HEADER_BG = '005293'
ROW_ALT_BG = 'F2F7FC'

def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, border_color='CCCCCC', border_size='4'):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def create_header_footer(doc):
    """创建页眉页脚"""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run('BMS 博客管理系统 - 详细设计说明书')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('第 ')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    # 页码
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = footer_para.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run3 = footer_para.add_run(' 页')
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(128, 128, 128)

def add_styled_heading(doc, text, level=1):
    """添加样式化标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = PRIMARY_COLOR
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def add_paragraph_with_style(doc, text, first_line_indent=True, bold=False):
    """添加格式化段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def create_data_table(doc, headers, data, first_col_highlight=False):
    """创建数据表格"""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_row = table.rows[0]
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
        set_cell_shading(header_row.cells[i], HEADER_BG)
        set_cell_border(header_row.cells[i], HEADER_BG)
        for para in header_row.cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        bg_color = ROW_ALT_BG if i % 2 == 0 else 'FFFFFF'
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = str(cell_data)
            if j == 0 and first_col_highlight:
                set_cell_shading(row.cells[j], 'F0F7FC')
            else:
                set_cell_shading(row.cells[j], bg_color)
            set_cell_border(row.cells[j])
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    return table

def add_code_block(doc, code_text, language=''):
    """添加代码块"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def create_report():
    """创建详细设计说明书"""
    doc = Document()

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 创建页眉页脚
    create_header_footer(doc)

    # ========== 封面 ==========
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('BMS 博客管理系统')
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('详细设计说明书')
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('━' * 30)
    run.font.color.rgb = ACCENT_COLOR
    run.font.size = Pt(14)

    en_title = doc.add_paragraph()
    en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = en_title.add_run('Blog Management System\nDetailed Design Document')
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = 'Arial'

    for _ in range(6):
        doc.add_paragraph()

    version_data = [
        ('文档版本', 'V1.0.0'),
        ('对应版本', 'V1.0.0'),
        ('编制日期', '2026-03-26'),
        ('编制人员', 'Sisyphus AI Agent'),
        ('技术框架', 'Snowy v3.6.1'),
        ('密级', '内部文档')
    ]

    version_table = doc.add_table(rows=len(version_data), cols=2)
    version_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, value) in enumerate(version_data):
        row = version_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(12)
        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = PRIMARY_COLOR

    doc.add_page_break()

    # ========== 文档修订历史 ==========
    add_styled_heading(doc, '文档修订历史', level=1)

    headers = ['版本号', '修订日期', '修订内容', '修订人', '审核人']
    data = [
        ('V1.0.0', '2026-03-26', '初始版本，完成详细设计说明', 'Sisyphus', '-'),
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()
    doc.add_page_break()

    # ========== 目录 ==========
    add_styled_heading(doc, '目  录', level=1)

    toc_items = [
        ('1', '引言', '4'),
        ('1.1', '编写目的', '4'),
        ('1.2', '项目背景', '4'),
        ('1.3', '定义与缩写', '4'),
        ('2', '系统架构设计', '5'),
        ('2.1', '系统总体架构', '5'),
        ('2.2', '技术架构', '5'),
        ('2.3', '部署架构', '6'),
        ('3', '模块设计', '7'),
        ('3.1', '文章管理模块', '7'),
        ('3.2', '分类管理模块', '8'),
        ('3.3', '标签管理模块', '9'),
        ('3.4', '评论管理模块', '10'),
        ('3.5', '媒体库模块', '11'),
        ('4', '数据库设计', '12'),
        ('4.1', '数据库概述', '12'),
        ('4.2', '表结构设计', '12'),
        ('4.3', '索引设计', '14'),
        ('5', '接口设计', '15'),
        ('5.1', '接口规范', '15'),
        ('5.2', '文章管理接口', '15'),
        ('5.3', '分类管理接口', '16'),
        ('5.4', '标签管理接口', '17'),
        ('5.5', '评论管理接口', '17'),
        ('5.6', '媒体库接口', '18'),
        ('6', '安全设计', '19'),
        ('6.1', '认证与授权', '19'),
        ('6.2', '数据安全', '19'),
        ('6.3', '接口安全', '20'),
        ('7', '性能设计', '21'),
        ('7.1', '性能指标', '21'),
        ('7.2', '缓存策略', '21'),
        ('7.3', '数据库优化', '22'),
        ('8', '附录', '23'),
    ]

    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    for i, (level, title, page) in enumerate(toc_items):
        row = toc_table.rows[i]
        row.cells[0].text = level
        row.cells[1].text = title
        row.cells[2].text = page
        for j, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(11)
        if not '.' in level:
            for run in row.cells[0].paragraphs[0].runs:
                run.font.bold = True
            for run in row.cells[1].paragraphs[0].runs:
                run.font.bold = True
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(10)
        row.cells[2].width = Cm(1.5)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()

    # ========== 1. 引言 ==========
    add_styled_heading(doc, '1. 引言', level=1)

    add_styled_heading(doc, '1.1 编写目的', level=2)
    add_paragraph_with_style(doc, '本文档是BMS博客管理系统的详细设计说明书，旨在为系统的开发、测试和维护提供技术参考。文档详细描述了系统的架构设计、模块设计、数据库设计、接口设计和安全设计等内容，是开发人员实现系统功能的主要依据。')

    add_paragraph_with_style(doc, '本文档的预期读者包括：', first_line_indent=False)
    p = doc.add_paragraph('• 系统开发人员：作为系统实现的参考指南')
    p.paragraph_format.left_indent = Cm(0.5)
    p = doc.add_paragraph('• 测试人员：作为测试用例设计的依据')
    p.paragraph_format.left_indent = Cm(0.5)
    p = doc.add_paragraph('• 运维人员：作为系统部署和维护的参考')
    p.paragraph_format.left_indent = Cm(0.5)
    p = doc.add_paragraph('• 项目管理人员：作为项目进度和质量控制的依据')
    p.paragraph_format.left_indent = Cm(0.5)

    add_styled_heading(doc, '1.2 项目背景', level=2)
    add_paragraph_with_style(doc, 'BMS博客管理系统是一个基于Snowy v3.6.1框架开发的轻量级内容管理平台，旨在为个人博主、内容创作者和小型媒体团队提供专业、易用的博客管理解决方案。系统采用前后端分离架构，后端基于Spring Boot 3.x，前端基于Vue 3 + Ant Design Vue。')

    add_styled_heading(doc, '1.3 定义与缩写', level=2)

    headers = ['术语/缩写', '定义']
    data = [
        ('BMS', 'Blog Management System，博客管理系统'),
        ('API', 'Application Programming Interface，应用程序接口'),
        ('REST', 'Representational State Transfer，表述性状态转移'),
        ('JWT', 'JSON Web Token，JSON网络令牌'),
        ('CRUD', 'Create, Read, Update, Delete，增删改查'),
        ('ORM', 'Object Relational Mapping，对象关系映射'),
        ('DTO', 'Data Transfer Object，数据传输对象'),
        ('VO', 'View Object，视图对象')
    ]
    create_data_table(doc, headers, data, first_col_highlight=True)
    doc.add_paragraph()

    # ========== 2. 系统架构设计 ==========
    add_styled_heading(doc, '2. 系统架构设计', level=1)

    add_styled_heading(doc, '2.1 系统总体架构', level=2)
    add_paragraph_with_style(doc, '系统采用经典的三层架构模式，结合Snowy框架的插件化设计理念，将业务模块以插件形式组织，实现模块间的松耦合。整体架构分为表现层、业务逻辑层和数据访问层。')

    # 架构层次表格
    headers = ['层次', '组件', '职责']
    data = [
        ('表现层', 'Vue 3 + Ant Design Vue', '用户界面展示、用户交互、表单验证'),
        ('接口层', 'Spring MVC Controller', 'RESTful API暴露、参数校验、权限控制'),
        ('业务层', 'Service', '业务逻辑处理、事务管理、数据组装'),
        ('数据层', 'MyBatis-Plus Mapper', '数据持久化、SQL执行、缓存管理'),
        ('基础设施', 'Redis、MySQL、MinIO', '数据存储、缓存服务、文件存储')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '2.2 技术架构', level=2)

    headers = ['层级', '技术选型', '版本', '说明']
    data = [
        ('前端框架', 'Vue', '3.x', '渐进式JavaScript框架'),
        ('UI组件库', 'Ant Design Vue', '4.x', '企业级UI组件库'),
        ('构建工具', 'Vite', '5.x', '下一代前端构建工具'),
        ('状态管理', 'Pinia', '2.x', 'Vue 3官方状态管理'),
        ('后端框架', 'Spring Boot', '3.x', 'Java企业级框架'),
        ('ORM框架', 'MyBatis-Plus', '3.x', '增强版MyBatis'),
        ('权限框架', 'Sa-Token', '1.x', '轻量级权限框架'),
        ('数据库', 'MySQL', '8.0+', '关系型数据库'),
        ('缓存', 'Redis', '6.0+', '内存数据库'),
        ('API文档', 'Knife4j', '4.x', 'Swagger增强UI')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '2.3 部署架构', level=2)
    add_paragraph_with_style(doc, '系统支持单机部署和集群部署两种模式。生产环境建议采用Nginx反向代理 + 应用集群 + 数据库主从的部署方案。')

    # 部署架构说明
    deploy_text = """
┌─────────────────────────────────────────────────────────────┐
│                         用户请求                              │
└─────────────────────────────────────────────────────────────┘
                              ↓
┌─────────────────────────────────────────────────────────────┐
│                    Nginx (反向代理/负载均衡)                   │
│                    监听端口: 80/443                           │
└─────────────────────────────────────────────────────────────┘
                    ↙                      ↘
        ┌──────────────────┐      ┌──────────────────┐
        │   前端静态资源     │      │   后端API服务     │
        │   (Vue 3应用)     │      │   (Spring Boot)   │
        │   端口: 81        │      │   端口: 82        │
        └──────────────────┘      └──────────────────┘
                                          ↓
                          ┌───────────────────────────┐
                          │      Redis 缓存集群        │
                          │      端口: 6379           │
                          └───────────────────────────┘
                                          ↓
                          ┌───────────────────────────┐
                          │     MySQL 主从集群         │
                          │     端口: 3306           │
                          └───────────────────────────┘
"""
    add_code_block(doc, deploy_text)

    # ========== 3. 模块设计 ==========
    add_styled_heading(doc, '3. 模块设计', level=1)

    add_styled_heading(doc, '3.1 文章管理模块', level=2)

    add_styled_heading(doc, '3.1.1 模块概述', level=3)
    add_paragraph_with_style(doc, '文章管理模块是系统的核心模块，负责博客文章的全生命周期管理，包括文章的创建、编辑、发布、撤回、删除等功能。模块支持富文本编辑、Markdown格式、文章分类标签关联、定时发布等特性。')

    add_styled_heading(doc, '3.1.2 功能清单', level=3)

    headers = ['功能编号', '功能名称', '功能描述', '优先级']
    data = [
        ('F-ART-001', '文章发布', '创建并发布博客文章', 'P1'),
        ('F-ART-002', '文章编辑', '修改文章内容，保存版本历史', 'P1'),
        ('F-ART-003', '文章删除', '逻辑删除文章，支持批量', 'P1'),
        ('F-ART-004', '文章列表', '分页查询，支持筛选排序', 'P1'),
        ('F-ART-005', '文章详情', '查看文章完整信息', 'P1'),
        ('F-ART-006', '发布/撤回', '控制文章发布状态', 'P1'),
        ('F-ART-007', '置顶/推荐', '设置文章置顶或推荐', 'P2'),
        ('F-ART-008', '定时发布', '设置定时发布时间', 'P2')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '3.1.3 类设计', level=3)

    class_text = """
// 文章实体类
@Entity
public class BmsArticle {
    @TableId
    private String id;              // 主键ID
    private String title;           // 文章标题
    private String summary;         // 文章摘要
    private String content;         // 文章内容(HTML/Markdown)
    private String coverImage;      // 封面图片URL
    private String categoryId;      // 分类ID
    private String authorId;        // 作者ID
    private String status;          // 状态: DRAFT/PUBLISHED
    private Integer viewCount;      // 阅读量
    private Integer likeCount;      // 点赞数
    private Integer commentCount;   // 评论数
    private Integer isTop;          // 是否置顶
    private Integer isRecommend;    // 是否推荐
    private DateTime publishTime;   // 发布时间
    private DateTime scheduledTime; // 定时发布时间
    // ... 公共字段省略
}

// 文章服务接口
public interface BmsArticleService {
    CommonResult<PageResult<BmsArticle>> page(BmsArticlePageParam param);
    CommonResult<BmsArticle> detail(String id);
    CommonResult<String> add(BmsArticleAddParam param);
    CommonResult<String> edit(BmsArticleEditParam param);
    CommonResult<String> delete(List<String> idList);
    CommonResult<String> publish(String id);
    CommonResult<String> unpublish(String id);
}
"""
    add_code_block(doc, class_text)

    add_styled_heading(doc, '3.2 分类管理模块', level=2)

    add_styled_heading(doc, '3.2.1 模块概述', level=3)
    add_paragraph_with_style(doc, '分类管理模块负责文章分类的层级化管理，支持多级分类结构，提供分类的增删改查、排序、启用禁用等功能。分类采用树形结构存储，支持无限层级扩展。')

    add_styled_heading(doc, '3.2.2 功能清单', level=3)

    headers = ['功能编号', '功能名称', '功能描述', '优先级']
    data = [
        ('F-CAT-001', '分类新增', '创建新的文章分类', 'P1'),
        ('F-CAT-002', '分类编辑', '修改分类名称、排序等', 'P1'),
        ('F-CAT-003', '分类删除', '删除分类（校验关联）', 'P1'),
        ('F-CAT-004', '分类树', '获取树形分类结构', 'P1'),
        ('F-CAT-005', '分类启用/禁用', '控制分类可用状态', 'P2')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '3.2.3 数据结构', level=3)

    category_text = """
// 分类实体类
@Entity
public class BmsCategory {
    @TableId
    private String id;          // 主键ID
    private String name;        // 分类名称
    private String code;        // 分类编码(唯一)
    private String parentId;    // 父分类ID (0表示根节点)
    private Integer sortCode;   // 排序码
    private String status;      // 状态: ENABLE/DISABLE
    private String remark;      // 备注
    // ... 公共字段省略
}
"""
    add_code_block(doc, category_text)

    add_styled_heading(doc, '3.3 标签管理模块', level=2)

    add_styled_heading(doc, '3.3.1 模块概述', level=3)
    add_paragraph_with_style(doc, '标签管理模块提供文章标签的维护功能，标签与文章为多对多关系。模块支持标签的颜色自定义、排序管理，可用于前台标签云展示。')

    add_styled_heading(doc, '3.3.2 功能清单', level=3)

    headers = ['功能编号', '功能名称', '功能描述', '优先级']
    data = [
        ('F-TAG-001', '标签新增', '创建新的文章标签', 'P1'),
        ('F-TAG-002', '标签编辑', '修改标签名称、颜色', 'P1'),
        ('F-TAG-003', '标签删除', '删除标签（解除关联）', 'P1'),
        ('F-TAG-004', '标签列表', '获取所有标签', 'P1'),
        ('F-TAG-005', '标签分页', '分页查询标签', 'P2')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '3.4 评论管理模块', level=2)

    add_styled_heading(doc, '3.4.1 模块概述', level=3)
    add_paragraph_with_style(doc, '评论管理模块负责访客评论的管理，支持评论审核机制、嵌套回复、敏感词过滤等功能。评论需要管理员审核通过后才能在前台显示。')

    add_styled_heading(doc, '3.4.2 功能清单', level=3)

    headers = ['功能编号', '功能名称', '功能描述', '优先级']
    data = [
        ('F-CMT-001', '发表评论', '访客提交评论', 'P1'),
        ('F-CMT-002', '评论审核', '管理员审核评论', 'P1'),
        ('F-CMT-003', '评论回复', '回复已有评论', 'P2'),
        ('F-CMT-004', '评论删除', '删除不当评论', 'P1'),
        ('F-CMT-005', '评论列表', '分页查询评论', 'P1'),
        ('F-CMT-006', '敏感词过滤', '自动过滤敏感内容', 'P2')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '3.4.3 状态流转', level=3)
    add_paragraph_with_style(doc, '评论状态包括：PENDING(待审核)、APPROVED(已通过)、REJECTED(已拒绝)。状态流转规则：')

    state_text = """
┌─────────┐    审核通过    ┌──────────┐
│ PENDING │ ─────────────→│ APPROVED │
│ (待审核) │               │ (已通过)  │
└─────────┘               └──────────┘
     │
     │ 审核拒绝
     ↓
┌──────────┐
│ REJECTED │
│ (已拒绝)  │
└──────────┘
"""
    add_code_block(doc, state_text)

    add_styled_heading(doc, '3.5 媒体库模块', level=2)

    add_styled_heading(doc, '3.5.1 模块概述', level=3)
    add_paragraph_with_style(doc, '媒体库模块负责系统媒体资源的管理，支持图片、视频、附件等多种类型文件的上传、预览、删除等操作。文件存储采用MinIO或本地存储方案。')

    add_styled_heading(doc, '3.5.2 功能清单', level=3)

    headers = ['功能编号', '功能名称', '功能描述', '优先级']
    data = [
        ('F-MED-001', '文件上传', '上传媒体文件', 'P1'),
        ('F-MED-002', '文件列表', '分页查询媒体', 'P1'),
        ('F-MED-003', '文件预览', '预览媒体内容', 'P1'),
        ('F-MED-004', '文件删除', '删除媒体文件', 'P1'),
        ('F-MED-005', '类型筛选', '按文件类型筛选', 'P2')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    # ========== 4. 数据库设计 ==========
    add_styled_heading(doc, '4. 数据库设计', level=1)

    add_styled_heading(doc, '4.1 数据库概述', level=2)
    add_paragraph_with_style(doc, '系统采用MySQL 8.0作为主数据库，遵循Snowy框架的数据库设计规范。数据库名称：bms，字符集：utf8mb4，排序规则：utf8mb4_general_ci。')

    add_styled_heading(doc, '4.2 表结构设计', level=2)

    add_styled_heading(doc, '4.2.1 表清单', level=3)

    headers = ['序号', '表名', '中文名称', '预估数据量']
    data = [
        ('1', 'BMS_ARTICLE', '文章表', '万级'),
        ('2', 'BMS_CATEGORY', '分类表', '百级'),
        ('3', 'BMS_TAG', '标签表', '百级'),
        ('4', 'BMS_ARTICLE_TAG', '文章标签关联表', '十万级'),
        ('5', 'BMS_COMMENT', '评论表', '十万级'),
        ('6', 'BMS_MEDIA', '媒体资源表', '万级'),
        ('7', 'BMS_ARTICLE_VERSION', '文章版本表', '十万级')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '4.2.2 文章表(BMS_ARTICLE)', level=3)

    headers = ['字段名', '类型', '可空', '说明']
    data = [
        ('ID', 'VARCHAR(32)', '否', '主键ID'),
        ('TITLE', 'VARCHAR(200)', '否', '文章标题'),
        ('SUMMARY', 'VARCHAR(500)', '是', '文章摘要'),
        ('CONTENT', 'LONGTEXT', '否', '文章内容'),
        ('COVER_IMAGE', 'VARCHAR(255)', '是', '封面图片URL'),
        ('CATEGORY_ID', 'VARCHAR(32)', '是', '分类ID'),
        ('AUTHOR_ID', 'VARCHAR(32)', '否', '作者ID'),
        ('STATUS', 'VARCHAR(20)', '否', '状态(DRAFT/PUBLISHED)'),
        ('VIEW_COUNT', 'INT', '否', '阅读量，默认0'),
        ('LIKE_COUNT', 'INT', '否', '点赞数，默认0'),
        ('COMMENT_COUNT', 'INT', '否', '评论数，默认0'),
        ('IS_TOP', 'TINYINT', '否', '是否置顶，默认0'),
        ('IS_RECOMMEND', 'TINYINT', '否', '是否推荐，默认0'),
        ('PUBLISH_TIME', 'DATETIME', '是', '发布时间'),
        ('SCHEDULED_TIME', 'DATETIME', '是', '定时发布时间'),
        ('CREATE_TIME', 'DATETIME', '否', '创建时间'),
        ('CREATE_USER', 'VARCHAR(32)', '否', '创建人'),
        ('UPDATE_TIME', 'DATETIME', '是', '更新时间'),
        ('UPDATE_USER', 'VARCHAR(32)', '是', '更新人'),
        ('DELETE_FLAG', 'TINYINT', '否', '删除标记，默认0')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '4.2.3 评论表(BMS_COMMENT)', level=3)

    headers = ['字段名', '类型', '可空', '说明']
    data = [
        ('ID', 'VARCHAR(32)', '否', '主键ID'),
        ('ARTICLE_ID', 'VARCHAR(32)', '否', '文章ID'),
        ('PARENT_ID', 'VARCHAR(32)', '是', '父评论ID'),
        ('NICKNAME', 'VARCHAR(50)', '否', '评论者昵称'),
        ('EMAIL', 'VARCHAR(100)', '是', '评论者邮箱'),
        ('CONTENT', 'VARCHAR(1000)', '否', '评论内容'),
        ('STATUS', 'VARCHAR(20)', '否', '状态(PENDING/APPROVED/REJECTED)'),
        ('REPLY_CONTENT', 'VARCHAR(500)', '是', '管理员回复内容'),
        ('CREATE_TIME', 'DATETIME', '否', '创建时间'),
        ('DELETE_FLAG', 'TINYINT', '否', '删除标记')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '4.3 索引设计', level=2)

    add_styled_heading(doc, '4.3.1 文章表索引', level=3)

    headers = ['索引名', '类型', '字段', '说明']
    data = [
        ('PRIMARY', '主键', 'ID', '主键索引'),
        ('IDX_ARTICLE_CATEGORY', '普通', 'CATEGORY_ID', '分类查询优化'),
        ('IDX_ARTICLE_STATUS', '普通', 'STATUS', '状态筛选优化'),
        ('IDX_ARTICLE_PUBLISH_TIME', '普通', 'PUBLISH_TIME', '时间排序优化'),
        ('IDX_ARTICLE_TOP', '组合', 'IS_TOP, PUBLISH_TIME', '置顶文章查询'),
        ('IDX_ARTICLE_AUTHOR', '普通', 'AUTHOR_ID', '作者文章查询')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    # ========== 5. 接口设计 ==========
    add_styled_heading(doc, '5. 接口设计', level=1)

    add_styled_heading(doc, '5.1 接口规范', level=2)

    add_styled_heading(doc, '5.1.1 RESTful API规范', level=3)

    headers = ['操作', 'HTTP方法', 'URL示例', '说明']
    data = [
        ('查询列表', 'GET', '/biz/article/page', '分页查询'),
        ('查询详情', 'GET', '/biz/article/detail?id=xxx', '单个查询'),
        ('新增', 'POST', '/biz/article/add', '创建资源'),
        ('修改', 'POST', '/biz/article/edit', '更新资源'),
        ('删除', 'POST', '/biz/article/delete', '批量删除')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '5.1.2 响应格式', level=3)

    response_text = """
// 统一响应格式
{
    "code": 200,           // 状态码：200成功，其他为失败
    "msg": "操作成功",      // 提示信息
    "data": { ... }        // 响应数据
}

// 分页响应格式
{
    "code": 200,
    "msg": "操作成功",
    "data": {
        "records": [...],  // 数据列表
        "total": 100,      // 总记录数
        "size": 10,        // 每页大小
        "current": 1,      // 当前页码
        "pages": 10        // 总页数
    }
}
"""
    add_code_block(doc, response_text)

    add_styled_heading(doc, '5.2 文章管理接口', level=2)

    headers = ['序号', '接口名称', '方法', '路径', '权限标识']
    data = [
        ('1', '获取文章分页', 'GET', '/biz/article/page', '/biz/article/page'),
        ('2', '获取文章列表', 'GET', '/biz/article/list', '/biz/article/list'),
        ('3', '获取文章详情', 'GET', '/biz/article/detail', '/biz/article/detail'),
        ('4', '添加文章', 'POST', '/biz/article/add', '/biz/article/add'),
        ('5', '编辑文章', 'POST', '/biz/article/edit', '/biz/article/edit'),
        ('6', '删除文章', 'POST', '/biz/article/delete', '/biz/article/delete'),
        ('7', '发布文章', 'POST', '/biz/article/publish', '/biz/article/publish'),
        ('8', '撤回文章', 'POST', '/biz/article/unpublish', '/biz/article/unpublish')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '5.3 分类管理接口', level=2)

    headers = ['序号', '接口名称', '方法', '路径', '权限标识']
    data = [
        ('1', '获取分类分页', 'GET', '/biz/category/page', '/biz/category/page'),
        ('2', '获取分类树', 'GET', '/biz/category/tree', '/biz/category/tree'),
        ('3', '获取分类详情', 'GET', '/biz/category/detail', '/biz/category/detail'),
        ('4', '添加分类', 'POST', '/biz/category/add', '/biz/category/add'),
        ('5', '编辑分类', 'POST', '/biz/category/edit', '/biz/category/edit'),
        ('6', '删除分类', 'POST', '/biz/category/delete', '/biz/category/delete'),
        ('7', '禁用分类', 'POST', '/biz/category/disableStatus', '/biz/category/disableStatus'),
        ('8', '启用分类', 'POST', '/biz/category/enableStatus', '/biz/category/enableStatus')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '5.4 标签管理接口', level=2)

    headers = ['序号', '接口名称', '方法', '路径', '权限标识']
    data = [
        ('1', '获取标签分页', 'GET', '/biz/tag/page', '/biz/tag/page'),
        ('2', '获取标签列表', 'GET', '/biz/tag/list', '/biz/tag/list'),
        ('3', '获取标签详情', 'GET', '/biz/tag/detail', '/biz/tag/detail'),
        ('4', '添加标签', 'POST', '/biz/tag/add', '/biz/tag/add'),
        ('5', '编辑标签', 'POST', '/biz/tag/edit', '/biz/tag/edit'),
        ('6', '删除标签', 'POST', '/biz/tag/delete', '/biz/tag/delete')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '5.5 评论管理接口', level=2)

    headers = ['序号', '接口名称', '方法', '路径', '权限标识']
    data = [
        ('1', '获取评论分页', 'GET', '/biz/comment/page', '/biz/comment/page'),
        ('2', '获取评论详情', 'GET', '/biz/comment/detail', '/biz/comment/detail'),
        ('3', '添加评论', 'POST', '/biz/comment/add', '/biz/comment/add'),
        ('4', '删除评论', 'POST', '/biz/comment/delete', '/biz/comment/delete'),
        ('5', '审核通过', 'POST', '/biz/comment/approve', '/biz/comment/approve'),
        ('6', '审核拒绝', 'POST', '/biz/comment/reject', '/biz/comment/reject')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '5.6 媒体库接口', level=2)

    headers = ['序号', '接口名称', '方法', '路径', '权限标识']
    data = [
        ('1', '获取媒体分页', 'GET', '/biz/media/page', '/biz/media/page'),
        ('2', '获取媒体详情', 'GET', '/biz/media/detail', '/biz/media/detail'),
        ('3', '添加媒体', 'POST', '/biz/media/add', '/biz/media/add'),
        ('4', '编辑媒体', 'POST', '/biz/media/edit', '/biz/media/edit'),
        ('5', '删除媒体', 'POST', '/biz/media/delete', '/biz/media/delete')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    # ========== 6. 安全设计 ==========
    add_styled_heading(doc, '6. 安全设计', level=1)

    add_styled_heading(doc, '6.1 认证与授权', level=2)

    add_styled_heading(doc, '6.1.1 认证机制', level=3)
    add_paragraph_with_style(doc, '系统采用Sa-Token框架实现认证功能，支持JWT Token认证方式。用户登录成功后，系统生成Token返回给客户端，客户端后续请求需在Header中携带Token。')

    auth_text = """
// 请求头格式
Authorization: Bearer {token}

// Token验证流程
1. 客户端携带Token请求接口
2. 服务端解析Token，验证有效性
3. 验证通过，获取用户信息，存入上下文
4. 验证失败，返回401未授权错误
"""
    add_code_block(doc, auth_text)

    add_styled_heading(doc, '6.1.2 授权机制', level=3)
    add_paragraph_with_style(doc, '系统采用RBAC(基于角色的访问控制)模型，通过用户-角色-权限三层结构实现细粒度的权限控制。每个接口都有对应的权限标识，用户必须拥有相应权限才能访问。')

    headers = ['权限类型', '说明', '示例']
    data = [
        ('菜单权限', '控制用户可见菜单', '文章管理菜单'),
        ('按钮权限', '控制页面按钮显示', '新增文章按钮'),
        ('接口权限', '控制API访问', '/biz/article/add'),
        ('数据权限', '控制数据范围', '仅查看自己文章')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '6.2 数据安全', level=2)

    add_styled_heading(doc, '6.2.1 敏感数据加密', level=3)

    headers = ['数据类型', '加密方式', '说明']
    data = [
        ('用户密码', 'SM3哈希', '国密算法，不可逆加密'),
        ('数据库密码', 'SM4加密', '配置文件加密存储'),
        ('Redis密码', 'SM4加密', '配置文件加密存储'),
        ('手机号', 'SM4加密', '数据库加密存储'),
        ('身份证号', 'SM4加密', '数据库加密存储')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '6.2.2 日志脱敏', level=3)
    add_paragraph_with_style(doc, '系统日志输出时，对敏感字段进行脱敏处理，防止敏感信息泄露。脱敏规则：手机号保留前3后4位，身份证号保留前4后4位，银行卡号保留后4位。')

    add_styled_heading(doc, '6.3 接口安全', level=2)

    headers = ['安全措施', '实现方式', '说明']
    data = [
        ('SQL注入防护', '参数化查询 + 白名单校验', 'MyBatis-Plus参数化'),
        ('XSS防护', '输入过滤 + 输出转义', '敏感标签过滤'),
        ('CSRF防护', 'Token验证', 'Sa-Token内置'),
        ('接口限流', 'Redis + 注解', '防止暴力请求'),
        ('请求签名', 'SM2签名验证', '可选，敏感接口')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    # ========== 7. 性能设计 ==========
    add_styled_heading(doc, '7. 性能设计', level=1)

    add_styled_heading(doc, '7.1 性能指标', level=2)

    headers = ['指标项', '目标值', '测试方法']
    data = [
        ('接口平均响应时间', '< 200ms', 'JMeter压测'),
        ('接口P95响应时间', '< 500ms', 'JMeter压测'),
        ('并发用户数', '≥ 500', 'JMeter压测'),
        ('数据库查询时间', '< 100ms', '慢查询日志'),
        ('缓存命中率', '> 80%', 'Redis监控'),
        ('系统可用性', '> 99.9%', '监控系统')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '7.2 缓存策略', level=2)

    add_styled_heading(doc, '7.2.1 缓存架构', level=3)
    add_paragraph_with_style(doc, '系统采用Redis作为缓存服务，支持本地缓存和分布式缓存两级缓存架构。热点数据优先走本地缓存，降低Redis访问压力。')

    cache_text = """
┌─────────────┐    命中    ┌─────────────┐
│  本地缓存    │ ←───────── │   请求处理   │
│  (Caffeine) │            │             │
└─────────────┘            └─────────────┘
      │ 未命中                    ↑
      ↓                          │
┌─────────────┐    命中          │
│ 分布式缓存   │ ────────────────┘
│   (Redis)   │
└─────────────┘
      │ 未命中
      ↓
┌─────────────┐
│   数据库     │
│   (MySQL)   │
└─────────────┘
"""
    add_code_block(doc, cache_text)

    add_styled_heading(doc, '7.2.2 缓存策略', level=3)

    headers = ['数据类型', '缓存方式', '过期时间', '更新策略']
    data = [
        ('分类树', 'Redis', '1小时', '定时刷新 + 主动更新'),
        ('标签列表', 'Redis', '1小时', '定时刷新 + 主动更新'),
        ('文章详情', 'Redis', '30分钟', 'LRU + 主动失效'),
        ('文章计数', 'Redis Hash', '实时', '定时同步到DB'),
        ('用户Session', 'Redis', '2小时', '滑动过期')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '7.3 数据库优化', level=2)

    add_styled_heading(doc, '7.3.1 查询优化', level=3)

    opt_items = [
        '• 避免SELECT *，只查询需要的字段',
        '• 合理使用索引，避免全表扫描',
        '• 批量查询替代循环查询',
        '• 分页查询添加LIMIT限制',
        '• 使用EXPLAIN分析慢查询'
    ]
    for item in opt_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    add_styled_heading(doc, '7.3.2 索引优化原则', level=3)

    idx_items = [
        '• 最左前缀原则：复合索引按查询条件顺序创建',
        '• 覆盖索引：将查询字段纳入索引，避免回表',
        '• 索引选择性：选择区分度高的字段建索引',
        '• 适度索引：避免过多索引影响写入性能'
    ]
    for item in idx_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()

    # ========== 8. 附录 ==========
    add_styled_heading(doc, '8. 附录', level=1)

    add_styled_heading(doc, '8.1 相关文档', level=2)

    headers = ['文档名称', '文档位置', '说明']
    data = [
        ('PRD文档', '.planning/docs/PRD.md', '产品需求文档'),
        ('功能清单', '.planning/docs/功能清单.md', '功能模块清单'),
        ('数据库设计', '.planning/docs/数据库设计.md', '数据库详细设计'),
        ('API接口清单', '.planning/docs/API接口测试清单.md', '接口详细说明'),
        ('测试报告', '.planning/docs/测试报告.md', '测试结果报告')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    add_styled_heading(doc, '8.2 环境配置', level=2)

    headers = ['组件', '版本', '配置要求']
    data = [
        ('JDK', '17+', 'JAVA_HOME配置'),
        ('MySQL', '8.0+', '字符集utf8mb4'),
        ('Redis', '6.0+', '密码认证'),
        ('Node.js', '18+', 'npm/yarn'),
        ('Maven', '3.6+', '镜像配置')
    ]
    create_data_table(doc, headers, data)
    doc.add_paragraph()

    # 文档结束
    doc.add_paragraph()
    end_line = doc.add_paragraph()
    end_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_line.add_run('—  文档结束  —')
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(12)

    # 保存文档
    output_path = 'BMS_详细设计说明书.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()